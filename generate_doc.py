from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)


def p(text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    para.paragraph_format.space_after = Pt(6)
    return para


# ═══ 标题 ═══
title = doc.add_paragraph()
title.alignment = 1  # center
run = title.add_run('乔信（聊天）功能实现说明')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(16)
run.bold = True

p('')

# ═══ 第一段：整体结构 ═══
p('一、功能概述', bold=True, size=12)
p('乔信是切瓜切菜小程序里的客户聊天模块，类似微信。'
  '代码由两个页面组成：chatWithCust（会话列表页）和 chatWithCustDetail（聊天详情页），'
  '位于 pages/module_others/pages/ 目录下。'
  '会话列表页底部有四个 Tab（叮客/通讯录/叮友/我），目前只有"叮客"（会话列表）和"我"（换头像、设昵称）可用。'
  '点击会话条目跳转 chatWithCustDetail 进入聊天界面，支持发送文字、图片、语音、视频四种消息。')

p('')

# ═══ 第二段：核心流程 ═══
p('二、消息收发流程', bold=True, size=12)
p('发送消息的核心函数是 chatWithCustDetail.js 里的 sendMsgExec(msg_type, msg, duration)。'
  '它调用 /qiaoxin/insertWechatMsgRecord 接口将消息写入数据库（wechat_type=3 表示客户间聊天），'
  '写入成功后调用 getNewData() 刷新聊天界面（调 /qiaoxin/selectWechatMsgRecord_StaffToCustomer_New 接口拉取新消息），'
  '同时调用 sendRabbitMqMsg() 通过 /evaluation/SendRabbitMqMsg 接口推送 RabbitMQ 通知给接收方。'
  '如果是图片/语音/视频消息，则需要先调用 wx.uploadFile 上传文件到阿里云 OSS（接口地址 app.globalData.UploadFile），'
  '拿到 OSS 返回的 URL 后再作为消息内容发送。'
  '点击发送按钮触发 btnSendRecord()，输入框回车触发 addSendRecord()，两者都调用 sendMsgExec(1, msg)。')

p('')
p('接收消息走的是实时推送：app.js 里的 chatWithCustDetail_GetNewMsg() 函数（约3002行）监听 WebSocket 推送，'
  '收到消息后自动调当前页面的 getNewData() 刷新，同时通知前置页面（会话列表）调 selectWechatMsgRecord_Last() 刷新最后一条消息。'
  'WebSocket 连接由 utils/RMQV3.js 管理，登录后自动连到 RabbitMQ，订阅个人队列 "{环境}_qgqc_{用户ID}"。')

p('')

# ═══ 第三段：接口和文件清单 ═══
p('三、关键接口与文件', bold=True, size=12)
p('接口（都在 https://mb.fsmbdlkj.com 下）：', bold=False)
p('  • /qiaoxin/insertWechatMsgRecord —— 发送消息（POST）')
p('  • /qiaoxin/selectWechatMsgRecord_StaffToCustomer_Old —— 拉历史消息，每页20条，max_id 控制分页（GET）')
p('  • /qiaoxin/selectWechatMsgRecord_StaffToCustomer_New —— 拉新消息，min_id 控制（GET）')
p('  • /qiaoxin/selectWeChatMsgRecord_QGQC_Customer —— 获取会话列表（GET）')
p('  • /qiaoxin/selectWechatMsgRecord_Last —— 获取某对话最后一条消息（GET）')
p('  • /evaluation/SendRabbitMqMsg —— RabbitMQ 推送通知（POST）')
p('  • /WX Restaurant/UploadFile —— 上传图片/视频/语音文件')

p('')
p('代码文件：')
p('  • pages/module_others/pages/chatWithCust/chatWithCust.js —— 会话列表页（initMsg 拉列表，selectWechatMsgRecord_Last 刷新单条）')
p('  • pages/module_others/pages/chatWithCustDetail/chatWithCustDetail.js —— 聊天详情页（sendMsgExec 发消息，getOldData/getNewData 拉消息，uploadImage/uploadVideo/uploadAudio 上传文件）')
p('  • utils/RMQV3.js —— RabbitMQ WebSocket 连接管理（initSocket 建连，订阅消息根据 msg.type 分发给 app.chatWithCustDetail_GetNewMsg）')
p('  • app.js（约3002行起）—— chatWithCustDetail_GetNewMsg() 收到推送后找到当前页面调刷新方法')

p('')
p('四、消息类型', bold=True, size=12)
p('msg_type=1 文字、msg_type=2 图片、msg_type=3 语音（含 duration 秒数）、msg_type=4 视频。'
  'chatWithCustDetail.wxml 里根据 sender_id 判断是自己还是对方消息，自己绿色气泡靠右，对方白色气泡靠左。')

# ── 保存 ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '乔信功能说明.docx')
doc.save(out)
print(f'Saved: {out}')
